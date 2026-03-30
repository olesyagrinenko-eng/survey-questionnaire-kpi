from __future__ import annotations

import os
import re
import tempfile
from typing import Any, List, Literal, Sequence

import pandas as pd
from docx import Document
from flask import Flask, after_this_request, flash, render_template, request, send_file

_BASE_DIR = os.path.dirname(os.path.abspath(__file__))
_DEFAULT_KPI = os.path.join(_BASE_DIR, "data", "KPI_framework_ads_FULL.xlsx")

AnswersMode = Literal["off", "all", "closed_only"]
AnswersStyle = Literal["bullets", "numbered", "inline"]

app = Flask(__name__)
app.secret_key = os.getenv("FLASK_SECRET_KEY", "dev-secret-key-change-me")
app.config["MAX_CONTENT_LENGTH"] = 16 * 1024 * 1024


def _kpi_excel_path() -> str:
    env = (os.getenv("KPI_FRAMEWORK_PATH") or "").strip()
    if env and os.path.isfile(env):
        return env
    return _DEFAULT_KPI


def _load_kpi_framework_df() -> pd.DataFrame:
    path = _kpi_excel_path()
    if not os.path.isfile(path):
        raise FileNotFoundError(f"Файл KPI не найден: {path}")
    df = pd.read_excel(path)
    df.columns = [str(c).strip() for c in df.columns]
    needed = [
        "Блок",
        "KPI",
        "Метка",
        "Формулировка вопроса",
        "Тип / шкала",
        "Варианты ответа (сокращенно)",
        "Обязательность",
        "Когда использовать",
    ]
    missing = [c for c in needed if c not in df.columns]
    if missing:
        raise ValueError(f"В KPI-файле нет колонок: {', '.join(missing)}")
    out = df[needed].copy()
    return out.reset_index(drop=True)


def _split_answer_options(raw_value: Any) -> List[str]:
    text = str(raw_value or "").strip()
    if not text or text == "-":
        return []
    parts = [p.strip() for p in re.split(r"\s*[/;]\s*", text) if p.strip()]
    return parts if parts else [text]


def _row_has_list_options(row: pd.Series) -> bool:
    return bool(_split_answer_options(row["Варианты ответа (сокращенно)"]))


def _answers_visible_for_row(row: pd.Series, answers_mode: AnswersMode) -> bool:
    if answers_mode == "off":
        return False
    if answers_mode == "all":
        return True
    return _row_has_list_options(row)


def _add_answer_paragraphs(
    doc: Document,
    options: List[str],
    style: AnswersStyle,
) -> None:
    if not options:
        return
    if style == "inline":
        joined = "; ".join(options)
        doc.add_paragraph(f"Варианты ответа: {joined}")
        return
    doc.add_paragraph("Варианты ответа:")
    if style == "numbered":
        for i, opt in enumerate(options, start=1):
            doc.add_paragraph(f"{i}. {opt}")
        return
    for opt in options:
        doc.add_paragraph(opt, style="List Bullet")


def _build_questionnaire_docx(
    df_filtered: pd.DataFrame,
    selected_labels_line: str,
    answers_mode: AnswersMode,
    answers_style: AnswersStyle,
    note_if_no_options: bool,
    include_rotation: bool,
    add_intro_instructions: bool,
) -> str:
    doc = Document()
    doc.add_heading("Анкета количественного опроса (по KPI Framework)", level=1)
    doc.add_paragraph("Документ сформирован автоматически (survey-questionnaire-kpi).")

    criteria = [
        f"Выбранные показатели: {selected_labels_line}",
        f"Варианты ответа в DOCX: {answers_mode} ({answers_style})",
        f"Пояснение для вопросов без списка в файле: {'да' if note_if_no_options else 'нет'}",
        f"Ротация вариантов (single/multi): {'да' if include_rotation else 'нет'}",
    ]
    doc.add_paragraph("Параметры сборки:\n- " + "\n- ".join(criteria))

    if add_intro_instructions:
        doc.add_heading("Общие инструкции интервьюеру", level=2)
        doc.add_paragraph(
            "Показывайте креатив перед блоком оценки. Соблюдайте скрининг, "
            "не подсказывайте ответы, фиксируйте ответ дословно там, где вопрос открытый."
        )

    for idx, row in df_filtered.reset_index(drop=True).iterrows():
        label = str(row["Метка"]).strip()
        kpi_name = str(row["KPI"]).strip()
        block = str(row["Блок"]).strip()
        question = str(row["Формулировка вопроса"]).strip()
        q_type = str(row["Тип / шкала"]).strip()
        obligatory = str(row["Обязательность"]).strip()
        when_use = str(row["Когда использовать"]).strip()

        doc.add_heading(f"Q{idx + 1}. [{label}] {kpi_name}", level=2)
        doc.add_paragraph(f"Блок: {block}")
        doc.add_paragraph(f"Формулировка: {question}")
        doc.add_paragraph(f"Тип вопроса: {q_type}")
        doc.add_paragraph(f"Обязательность: {obligatory}")
        doc.add_paragraph(f"Условие показа: {when_use}")

        options = _split_answer_options(row["Варианты ответа (сокращенно)"])
        if _answers_visible_for_row(row, answers_mode):
            if options:
                _add_answer_paragraphs(doc, options, answers_style)
            elif note_if_no_options and answers_mode == "all":
                doc.add_paragraph(
                    "Варианты ответа: в таблице KPI не заданы явно — используйте шкалу из «Тип / шкала» "
                    "или сформулируйте список вручную."
                )

        if (
            include_rotation
            and answers_mode != "off"
            and options
            and any(token in q_type.lower() for token in ["multi", "single"])
        ):
            doc.add_paragraph("Ротация: ротировать порядок вариантов ответа.")

    fd, output_path = tempfile.mkstemp(suffix=".docx", prefix="questionnaire_")
    os.close(fd)
    doc.save(output_path)
    return output_path


def _framework_items(framework_df: pd.DataFrame) -> List[dict]:
    items: List[dict] = []
    for i, row in framework_df.iterrows():
        label = str(row["Метка"]).strip()
        kpi_name = str(row["KPI"]).strip()
        formulation = str(row["Формулировка вопроса"]).strip()
        block = str(row["Блок"]).strip()
        has_opts = _row_has_list_options(row)
        search_blob = f"{label} {kpi_name} {formulation} {block}".lower()
        items.append(
            {
                "id": int(i),
                "block": block,
                "kpi": kpi_name,
                "label": label,
                "formulation": formulation,
                "q_type": str(row["Тип / шкала"]).strip(),
                "has_options": has_opts,
                "search_blob": search_blob,
            }
        )
    return items


def _parse_selected_indices(form: Any) -> List[int]:
    raw: Sequence[str] = form.getlist("q")
    out: List[int] = []
    for x in raw:
        try:
            out.append(int(x))
        except (TypeError, ValueError):
            continue
    return out


def _template_ctx(
    *,
    framework_df: pd.DataFrame,
    items: List[dict],
    selected_ids: List[int],
    answers_mode: str,
    answers_style: str,
    note_if_no_options: bool,
    include_rotation: bool,
    add_intro_instructions: bool,
    when_filter: str,
    kpi_path: str,
) -> dict:
    sel_set = {int(x) for x in selected_ids}
    return {
        "items": items,
        "selected_ids": sel_set,
        "answers_mode": answers_mode,
        "answers_style": answers_style,
        "note_if_no_options": note_if_no_options,
        "include_rotation": include_rotation,
        "add_intro_instructions": add_intro_instructions,
        "when_filter": when_filter,
        "total_questions": int(framework_df.shape[0]),
        "kpi_source_label": kpi_path,
    }


@app.route("/", methods=["GET", "POST"])
def index():
    kpi_path = _kpi_excel_path()
    try:
        framework_df = _load_kpi_framework_df()
    except Exception as exc:
        flash(f"Не удалось прочитать KPI framework: {exc}", "error")
        return render_template(
            "index.html",
            **_template_ctx(
                framework_df=pd.DataFrame(),
                items=[],
                selected_ids=[],
                answers_mode="all",
                answers_style="bullets",
                note_if_no_options=True,
                include_rotation=True,
                add_intro_instructions=True,
                when_filter="",
                kpi_path=kpi_path,
            ),
        )

    items = _framework_items(framework_df)

    if request.method == "GET":
        return render_template(
            "index.html",
            **_template_ctx(
                framework_df=framework_df,
                items=items,
                selected_ids=[it["id"] for it in items],
                answers_mode="all",
                answers_style="bullets",
                note_if_no_options=True,
                include_rotation=True,
                add_intro_instructions=True,
                when_filter="",
                kpi_path=kpi_path,
            ),
        )

    answers_mode = (request.form.get("answers_mode") or "all").strip()
    if answers_mode not in ("off", "all", "closed_only"):
        answers_mode = "all"
    answers_style = (request.form.get("answers_style") or "bullets").strip()
    if answers_style not in ("bullets", "numbered", "inline"):
        answers_style = "bullets"

    note_if_no_options = bool(request.form.get("note_if_no_options"))
    include_rotation = bool(request.form.get("include_rotation"))
    add_intro_instructions = bool(request.form.get("add_intro_instructions"))
    when_filter = (request.form.get("when_filter") or "").strip()

    selected_ids = _parse_selected_indices(request.form)
    if not selected_ids:
        flash("Отметьте хотя бы один показатель (вопрос).", "error")
        return render_template(
            "index.html",
            **_template_ctx(
                framework_df=framework_df,
                items=items,
                selected_ids=[],
                answers_mode=answers_mode,
                answers_style=answers_style,
                note_if_no_options=note_if_no_options,
                include_rotation=include_rotation,
                add_intro_instructions=add_intro_instructions,
                when_filter=when_filter,
                kpi_path=kpi_path,
            ),
        )

    mask = framework_df.index.isin(selected_ids)
    filtered = framework_df.loc[mask].copy()
    order_map = {rid: pos for pos, rid in enumerate(selected_ids)}
    filtered["_sort"] = filtered.index.map(lambda i: order_map.get(int(i), 9999))
    filtered = filtered.sort_values("_sort").drop(columns="_sort")

    if when_filter:
        filtered = filtered[
            filtered["Когда использовать"].astype(str).str.contains(when_filter, case=False, na=False)
        ]

    if filtered.empty:
        flash("После фильтра «Когда использовать» не осталось ни одного вопроса.", "error")
        return render_template(
            "index.html",
            **_template_ctx(
                framework_df=framework_df,
                items=items,
                selected_ids=selected_ids,
                answers_mode=answers_mode,
                answers_style=answers_style,
                note_if_no_options=note_if_no_options,
                include_rotation=include_rotation,
                add_intro_instructions=add_intro_instructions,
                when_filter=when_filter,
                kpi_path=kpi_path,
            ),
        )

    labels_line = ", ".join(str(x) for x in filtered["Метка"].tolist())

    try:
        output_path = _build_questionnaire_docx(
            filtered,
            selected_labels_line=labels_line,
            answers_mode=answers_mode,  # type: ignore[arg-type]
            answers_style=answers_style,  # type: ignore[arg-type]
            note_if_no_options=note_if_no_options,
            include_rotation=include_rotation,
            add_intro_instructions=add_intro_instructions,
        )
    except Exception as exc:
        flash(f"Ошибка генерации DOCX: {exc}", "error")
        return render_template(
            "index.html",
            **_template_ctx(
                framework_df=framework_df,
                items=items,
                selected_ids=selected_ids,
                answers_mode=answers_mode,
                answers_style=answers_style,
                note_if_no_options=note_if_no_options,
                include_rotation=include_rotation,
                add_intro_instructions=add_intro_instructions,
                when_filter=when_filter,
                kpi_path=kpi_path,
            ),
        )

    @after_this_request
    def _cleanup_doc(response):  # noqa: ANN001
        try:
            os.remove(output_path)
        except OSError:
            pass
        return response

    return send_file(
        output_path,
        as_attachment=True,
        download_name="questionnaire_from_kpi_framework.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


if __name__ == "__main__":
    port = int(os.getenv("PORT", "5000"))
    app.run(host="0.0.0.0", port=port, debug=True)
