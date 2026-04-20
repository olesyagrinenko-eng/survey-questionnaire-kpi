# -*- coding: utf-8 -*-
"""Выгрузка структуры анкеты в .docx."""
from __future__ import annotations

import io
from typing import Any

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt


def spec_to_docx(spec: dict[str, Any]) -> bytes:
    meta = spec.get("meta") or {}
    blocks = spec.get("blocks") or []

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = meta.get("project_name") or "Анкета"
    h0 = doc.add_heading(title, 0)
    h0.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    p = doc.add_paragraph()
    p.add_run("Этап: ").bold = True
    p.add_run(meta.get("phase_label") or meta.get("phase") or "—")

    if meta.get("counts"):
        doc.add_paragraph()
        ph = doc.add_paragraph()
        ph.add_run("Материалы (количество): ").bold = True
        parts = []
        for k, v in meta["counts"].items():
            if int(v or 0) > 0:
                parts.append(f"{k}: {v}")
        ph.add_run(", ".join(parts) if parts else "—")

    assets = meta.get("stimulus_assets")
    if isinstance(assets, dict) and assets:
        doc.add_paragraph()
        ah = doc.add_paragraph()
        ah.add_run("Ссылки на стимулы (из конструктора): ").bold = True
        for key, rows in assets.items():
            if not isinstance(rows, list):
                continue
            for i, row in enumerate(rows, start=1):
                if isinstance(row, dict):
                    url = (row.get("url") or "").strip()
                    lab = (row.get("label") or "").strip()
                else:
                    url = str(row or "").strip()
                    lab = ""
                if not url:
                    continue
                line = f"{key} #{i}: {url}"
                if lab:
                    line += f" ({lab})"
                doc.add_paragraph(line, style="List Bullet")

    cn = (meta.get("client_notes") or "").strip()
    if cn:
        doc.add_paragraph()
        cn_p = doc.add_paragraph()
        cn_p.add_run("Пожелания заказчика: ").bold = True
        cn_p.add_run(cn)

    doc.add_paragraph()

    for block in blocks:
        doc.add_heading(block.get("title") or "Блок", level=1)
        instr = (block.get("programmer_instructions") or "").strip()
        if instr:
            pi = doc.add_paragraph()
            pi.add_run("Инструкция программисту: ").bold = True
            pi.add_run(instr)

        for q in block.get("questions") or []:
            doc.add_paragraph()
            qh = doc.add_paragraph()
            qh.add_run(f"[{q.get('id', '?')}] ").bold = True
            qh.add_run(q.get("text") or "")

            qt = doc.add_paragraph()
            qt.add_run("Тип вопроса: ").bold = True
            qt.add_run(_qtype_ru(q.get("qtype")))

            note = (q.get("programmer_note") or "").strip()
            if note:
                pn = doc.add_paragraph()
                pn.add_run("Заметка: ").italic = True
                pn.add_run(note)

            opts = q.get("options")
            if opts:
                ol = doc.add_paragraph()
                ol.add_run("Варианты ответа:").bold = True
                for o in opts:
                    doc.add_paragraph(str(o), style="List Bullet")

            anch = q.get("anchors")
            if anch:
                an = doc.add_paragraph()
                an.add_run("Шкала: ").bold = True
                an.add_run(
                    " — ".join(f"{k}: {v}" for k, v in sorted(anch.items(), key=lambda x: x[0]))
                )

            st = q.get("stimulus")
            if st:
                sg = doc.add_paragraph()
                sg.add_run("Привязка к стимулу: ").bold = True
                sg.add_run(f"{st.get('type')} #{st.get('index')}")
                au = (st.get("asset_url") or "").strip()
                if au:
                    al = doc.add_paragraph()
                    al.add_run("URL медиа: ").italic = True
                    al.add_run(au)

        doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def _qtype_ru(qtype: str | None) -> str:
    m = {
        "open": "Открытый ответ",
        "open_numeric": "Число (возраст и т.п.)",
        "single": "Один вариант из списка",
        "multi": "Несколько вариантов из списка",
        "multi_placeholder": "Множественный выбор (список задать в ТЗ)",
        "scale_1_9": "Шкала 1–9",
        "scale_1_5": "Шкала 1–5",
        "instruction": "Инструкция / экран без ответа",
        "click_map": "Клик-тест / карта внимания (координаты)",
        "yes_no": "Да / Нет",
    }
    return m.get(qtype or "", qtype or "—")
